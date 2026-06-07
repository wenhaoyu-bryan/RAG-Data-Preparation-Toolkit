"""Extract images and captions from Word documents."""

import os
import re
from typing import Dict, List
from docx import Document
from . import section_parser


def _is_caption_text(t: str) -> bool:
    if not t:
        return False
    if re.match(r'^(图表|图|Figure|Fig)[\s：:]*\d+', t.strip()):
        return True
    return len(t.strip()) <= 120


def extract_images_with_captions_from_docx(docx_path: str, output_dir: str = "extracted_images") -> Dict[str, Dict[str, str]]:
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc_name = os.path.splitext(os.path.basename(docx_path))[0]
    doc = Document(docx_path)
    image_info = {}
    image_counter = 1

    all_caption_paras = []
    for _ci, _p in enumerate(Document(docx_path).paragraphs):
        _t = _p.text.strip()
        if _is_caption_text(_t):
            all_caption_paras.append((_ci, _t))

    # Table images
    table_image_counter = 0
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    has_image = False
                    image_rel_id = None

                    for run in paragraph.runs:
                        for drawing in run._element.xpath('.//a:blip'):
                            embed_id = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed_id:
                                image_rel_id = embed_id
                                has_image = True
                                break
                        if not has_image:
                            for imagedata in run._element.xpath('.//*[local-name()="imagedata"]'):
                                embed_id = imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                if embed_id:
                                    image_rel_id = embed_id
                                    has_image = True
                                    break
                        if has_image:
                            break

                    if has_image and image_rel_id:
                        caption = ""
                        caption_index = None
                        if paragraph.text.strip():
                            caption = paragraph.text.strip()
                            caption_index = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}_para_{para_idx}"
                        if not caption and para_idx + 1 < len(cell.paragraphs):
                            next_para = cell.paragraphs[para_idx + 1]
                            if next_para.text.strip():
                                caption = next_para.text.strip()
                                caption_index = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}_para_{para_idx+1}"
                        if not caption and para_idx > 0:
                            prev_para = cell.paragraphs[para_idx - 1]
                            if prev_para.text.strip():
                                caption = prev_para.text.strip()
                                caption_index = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}_para_{para_idx-1}"

                        target_rel = None
                        for rel in doc.part.rels.values():
                            if rel.rId == image_rel_id:
                                target_rel = rel
                                break

                        if target_rel and "image" in target_rel.target_ref:
                            image_data = target_rel.target_part.blob
                            ext = _get_image_ext(target_rel.target_ref)
                            image_filename = f"{doc_name}_image_id____image{image_counter}{ext}"
                            image_path = os.path.join(output_dir, image_filename)

                            with open(image_path, 'wb') as f:
                                f.write(image_data)

                            image_info[f"image_{image_counter}"] = {
                                "filename": image_filename,
                                "caption": caption,
                                "para_index": f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}_para_{para_idx}",
                                "caption_index": caption_index,
                                "is_table_image": True,
                                "table_index": table_idx,
                            }
                            print(f"Extracted image: {image_filename}")
                            if caption:
                                print(f"  Caption: {caption}")
                            image_counter += 1
                            table_image_counter += 1

    print(f"Extracted {table_image_counter} images from tables")

    # Paragraph images
    for i, paragraph in enumerate(doc.paragraphs):
        image_rel_ids = []
        for run in paragraph.runs:
            for drawing in run._element.xpath('.//a:blip'):
                embed_id = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id:
                    image_rel_ids.append(embed_id)
            for imagedata in run._element.xpath('.//*[local-name()="imagedata"]'):
                embed_id = imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if embed_id:
                    image_rel_ids.append(embed_id)

        for image_rel_id in image_rel_ids:
            caption, caption_index = _find_caption(doc, paragraph, i)

            target_rel = None
            for rel in doc.part.rels.values():
                if rel.rId == image_rel_id:
                    target_rel = rel
                    break

            if target_rel and "image" in target_rel.target_ref:
                image_data = target_rel.target_part.blob
                ext = _get_image_ext(target_rel.target_ref)
                image_filename = f"{doc_name}_image_id____image{image_counter}{ext}"
                image_path = os.path.join(output_dir, image_filename)

                with open(image_path, 'wb') as f:
                    f.write(image_data)

                image_info[f"image_{image_counter}"] = {
                    "filename": image_filename,
                    "caption": caption,
                    "para_index": i,
                    "caption_index": caption_index,
                }
                print(f"Extracted image: {image_filename}")
                if caption:
                    print(f"  Caption: {caption}")
                image_counter += 1

    # Fallback caption matching
    _fallback_caption_match(image_info, all_caption_paras)

    return image_info


def _get_image_ext(target_ref: str) -> str:
    if target_ref.endswith('.png'):
        return '.png'
    elif target_ref.endswith('.jpg') or target_ref.endswith('.jpeg'):
        return '.jpg'
    elif target_ref.endswith('.gif'):
        return '.gif'
    elif target_ref.endswith('.bmp'):
        return '.bmp'
    return '.png'


def _find_caption(doc, paragraph, i: int):
    candidates = []
    cur_text = paragraph.text.strip()
    if cur_text:
        score = 2 if re.match(r'^(图表|图|Figure|Fig)[\s：:]*\d+', cur_text) else 1
        candidates.append((cur_text, i, score))

    for offset in range(1, 9):
        j = i + offset
        if j >= len(doc.paragraphs):
            break
        t = doc.paragraphs[j].text.strip()
        if not t:
            continue
        try:
            if section_parser.is_heading_paragraph(doc.paragraphs[j]):
                break
        except Exception:
            pass
        if _is_caption_text(t):
            score = 4 if re.match(r'^(图表|图|Figure|Fig)[\s：:]*\d+', t) else 2
            candidates.append((t, j, score))

    for offset in range(1, 9):
        j = i - offset
        if j < 0:
            break
        t = doc.paragraphs[j].text.strip()
        if not t:
            continue
        try:
            if section_parser.is_heading_paragraph(doc.paragraphs[j]):
                break
        except Exception:
            pass
        if _is_caption_text(t):
            score = 3 if re.match(r'^(图表|图|Figure|Fig)[\s：:]*\d+', t) else 1
            candidates.append((t, j, score))

    try:
        for run in paragraph.runs:
            for docpr in run._element.xpath('.//wp:docPr', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}):
                title = (docpr.get('title') or '').strip()
                descr = (docpr.get('descr') or '').strip()
                for t in (title, descr):
                    if t:
                        candidates.append((t, i, 4))
    except Exception:
        pass

    if candidates:
        candidates.sort(key=lambda x: (-x[2], abs(x[1] - i)))
        return candidates[0][0], candidates[0][1]

    caption = ""
    caption_index = None
    if i + 1 < len(doc.paragraphs):
        next_text = doc.paragraphs[i + 1].text.strip()
        if next_text:
            caption = next_text
            caption_index = i + 1
    if not caption and i > 0:
        prev_text = doc.paragraphs[i - 1].text.strip()
        if prev_text:
            caption = prev_text
            caption_index = i - 1
    return caption, caption_index


def _fallback_caption_match(image_info, all_caption_paras):
    if not image_info:
        return
    items = []
    for k, v in image_info.items():
        items.append((int(k.split('_')[1]), v.get('para_index', -1), k))
    items.sort()

    for _, para_idx, key in items:
        if image_info[key].get('caption'):
            continue
        best = None
        best_dist = 10 ** 9
        for (ci, txt) in all_caption_paras:
            if isinstance(para_idx, str) and para_idx.startswith('table_'):
                best = (ci, txt)
                break
            else:
                dist = abs(ci - para_idx)
                if ci <= para_idx:
                    dist -= 0.1
                if dist < best_dist:
                    best_dist = dist
                    best = (ci, txt)
        if best:
            image_info[key]['caption'] = best[1]
            image_info[key]['caption_index'] = best[0]


def create_enhanced_image_chunk_content(filename: str, target_section: str, caption: str) -> str:
    """Create a formatted image reference block for embedding in a chunk."""
    clean_caption = ""
    if caption:
        tmp = re.sub(r'^图表\s*\d+[\s：:]*', '', caption).strip()
        tmp = re.sub(r'\s+', ' ', tmp)
        clean_caption = f"Image content: {tmp}" if tmp else ""

    bracket = f"[Image: {filename}"
    if clean_caption:
        bracket += f"\n{clean_caption}"
    bracket += "]"
    lines = [bracket, f"Image location: {target_section}"]
    return "\n".join(lines)


def find_image_references_in_text(text: str) -> List[str]:
    image_patterns = [
        r'图\s*\d+', r'图片\s*\d+', r'附图\s*\d+',
        r'Figure\s*\d+', r'Fig\s*\d+',
        r'见下图', r'如图所示', r'参考图',
    ]
    references = []
    for pattern in image_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        references.extend(matches)
    return references


def identify_image_section(image_content: str, current_section_path: str) -> str:
    """Return the section path where an image should be attributed.

    Uses position-based attribution: the image belongs to whatever section
    was active at the image's location in the document.
    """
    if not current_section_path:
        return ""
    return current_section_path
